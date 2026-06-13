"""
Render a ResumeModel to a clean, ATS-friendly .docx (python-docx).

Word output is deliberately single-column plain text so applicant tracking
systems parse it perfectly, while keeping the navy/steel-blue accent styling.
"""

from __future__ import annotations

import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from schema import ResumeModel

HEADER = RGBColor(0x19, 0x19, 0x70)
ACCENT = RGBColor(0x46, 0x82, 0xB4)
TEXT = RGBColor(0x21, 0x21, 0x21)


def _set_base_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9)
    style.font.color.rgb = TEXT
    pf = style.paragraph_format
    pf.space_after = Pt(2)
    pf.space_before = Pt(0)


def _section(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = HEADER
    # bottom border (accent rule) under the heading
    pPr = p._p.get_or_add_pPr()
    borders = pPr.makeelement(qn("w:pBdr"), {})
    bottom = borders.makeelement(qn("w:bottom"), {
        qn("w:val"): "single", qn("w:sz"): "6", qn("w:space"): "1", qn("w:color"): "4682B4",
    })
    borders.append(bottom)
    pPr.append(borders)


def _bullets(doc: Document, items: list[str]) -> None:
    for b in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        p.add_run(b)


def render_docx(resume: ResumeModel) -> bytes:
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Pt(28)
        s.left_margin = s.right_margin = Pt(28)
    _set_base_font(doc)

    # Header
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = name_p.add_run(resume.name.upper())
    nr.bold = True
    nr.font.size = Pt(22)
    nr.font.color.rgb = HEADER

    if resume.title:
        t = doc.add_paragraph()
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tr = t.add_run(resume.title)
        tr.font.size = Pt(11)
        tr.font.color.rgb = ACCENT

    contact_bits = [b for b in [resume.email, resume.phone, resume.location,
                                resume.linkedin, resume.github] if b]
    if contact_bits:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.add_run("  |  ".join(contact_bits)).font.size = Pt(8.5)

    if resume.impact_metrics:
        band = doc.add_paragraph()
        band.alignment = WD_ALIGN_PARAGRAPH.CENTER
        br = band.add_run("   |   ".join(resume.impact_metrics))
        br.bold = True
        br.font.color.rgb = HEADER
        br.font.size = Pt(9)

    if resume.summary:
        _section(doc, "Executive Summary")
        doc.add_paragraph(resume.summary)

    if resume.competencies:
        _section(doc, "Core Technical Competencies")
        for comp in resume.competencies:
            p = doc.add_paragraph()
            r = p.add_run(f"{comp.category}: ")
            r.bold = True
            r.font.color.rgb = ACCENT
            p.add_run(comp.skills)

    if resume.experience:
        _section(doc, "Professional Experience")
        for x in resume.experience:
            head = doc.add_paragraph()
            hr = head.add_run(x.role)
            hr.bold = True
            hr.font.size = Pt(10)
            if x.dates:
                head.add_run(f"\t{x.dates}").bold = True
            sub = doc.add_paragraph()
            company = x.company + (f" | Client: {x.client}" if x.client else "")
            si = sub.add_run(company + (f"  —  {x.location}" if x.location else ""))
            si.italic = True
            si.font.size = Pt(8.8)
            if x.context:
                ctx = doc.add_paragraph()
                cr = ctx.add_run(x.context)
                cr.italic = True
                cr.font.size = Pt(8.6)
            for g in x.groups:
                if g.subhead:
                    gp = doc.add_paragraph()
                    gr = gp.add_run(g.subhead)
                    gr.bold = True
                    gr.font.color.rgb = ACCENT
                _bullets(doc, g.bullets)

    if resume.projects:
        _section(doc, "Key Projects")
        for p in resume.projects:
            if p.subhead:
                gp = doc.add_paragraph()
                gr = gp.add_run(p.subhead)
                gr.bold = True
                gr.font.color.rgb = ACCENT
            _bullets(doc, p.bullets)

    if resume.education:
        _section(doc, "Education")
        for e in resume.education:
            p = doc.add_paragraph()
            dr = p.add_run(e.degree)
            dr.bold = True
            if e.dates:
                p.add_run(f"\t{e.dates}").bold = True
            line2 = doc.add_paragraph()
            line2.add_run(e.institution + (f"   {e.score}" if e.score else ""))
            if e.coursework:
                cw = doc.add_paragraph()
                cwr = cw.add_run(f"Relevant Coursework: {e.coursework}")
                cwr.italic = True

    if resume.certifications:
        _section(doc, "Certifications & Professional Development")
        _bullets(doc, resume.certifications)

    if resume.languages or resume.target_roles:
        _section(doc, "Additional Information")
        if resume.languages:
            p = doc.add_paragraph()
            p.add_run("Languages: ").bold = True
            p.add_run(resume.languages)
        if resume.target_roles:
            p = doc.add_paragraph()
            p.add_run("Target Roles: ").bold = True
            p.add_run(resume.target_roles)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
